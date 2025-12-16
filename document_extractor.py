import io
import os
import base64
from typing import Dict, List, Any, Tuple
from PIL import Image
import pytesseract
import PyPDF2
from pdf2image import convert_from_bytes
from docx import Document
from openpyxl import load_workbook
from drill_qaqc_analyzer import DrillQAQCAnalyzer
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed

MAX_PARALLEL_PAGES = 5


class DocumentExtractor:
    
    @staticmethod
    def _get_openai_client():
        """Get OpenAI client using Replit AI Integrations or standard OpenAI"""
        api_key = os.environ.get('AI_INTEGRATIONS_OPENAI_API_KEY') or os.environ.get('OPENAI_API_KEY')
        
        if not api_key:
            raise ValueError("OpenAI API key not configured. Set AI_INTEGRATIONS_OPENAI_API_KEY or OPENAI_API_KEY environment variable.")
        
        base_url = os.environ.get('AI_INTEGRATIONS_OPENAI_BASE_URL')
        
        if base_url:
            return OpenAI(api_key=api_key, base_url=base_url)
        else:
            return OpenAI(api_key=api_key)
    
    @staticmethod
    def _extract_text_with_vision(image_bytes: bytes, context: str = "document") -> str:
        """Use GPT-4 Vision to extract text from images with better accuracy, especially for tables and structured data"""
        try:
            client = DocumentExtractor._get_openai_client()
            
            base64_image = base64.b64encode(image_bytes).decode('utf-8')
            
            response = client.chat.completions.create(
                model="gpt-5.1",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": f"""Extract ALL text from this {context} image with MAXIMUM ACCURACY. Pay special attention to:

CRITICAL - TABLES & STRUCTURED DATA:
- Resource/Reserve statements (tonnage, grade, contained metal by category: Measured, Indicated, Inferred, Proven, Probable)
- Financial tables (CAPEX, OPEX, NPV, IRR, cash flows)
- Drill hole data tables (collar, assay, survey data)
- Production schedules and metallurgical test results

FORMATTING RULES:
1. Preserve table structure using | delimiters (e.g., "Column1 | Column2 | Column3")
2. Keep numerical data EXACT (don't round or modify values)
3. Include units (tonnes, g/t, oz, %, meters, etc.)
4. Preserve headers and row labels
5. Include footnotes and notes

ALSO EXTRACT:
- All text, headings, and paragraphs
- Charts and diagram labels
- Technical specifications
- Legal disclaimers

Return ONLY the extracted text without commentary or analysis."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_completion_tokens=4096,
                reasoning_effort="high"
            )
            
            return response.choices[0].message.content.strip()
        except ValueError as ve:
            return f"VLM extraction unavailable: {str(ve)}"
        except Exception as e:
            return f"VLM extraction failed: {str(e)}"
    
    @staticmethod
    def _process_single_page(args: Tuple[bytes, int]) -> Tuple[int, str]:
        """Process a single page image and return (page_num, text)"""
        img_bytes, page_num = args
        page_text = DocumentExtractor._extract_text_with_vision(img_bytes, f"mining technical report page {page_num}")
        if page_text.startswith("VLM extraction failed") or page_text.startswith("VLM extraction unavailable"):
            return (page_num, "")
        return (page_num, f"\n\n========== PAGE {page_num} ==========\n{page_text}\n")
    
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Enhanced PDF extraction using GPT-4 Vision for maximum accuracy on tables and structured data"""
        try:
            # ALWAYS use VLM for technical/mining PDFs to ensure accurate table extraction
            # PyPDF2 is notoriously poor at extracting tables, resource statements, and structured data
            try:
                # Get total page count first
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                total_pages = len(pdf_reader.pages)
                
                # Process up to 500 pages with VLM for comprehensive extraction
                # Ensures complete coverage of NI 43-101 reports (typically 200-300 pages)
                # including resource estimates, drill results, economics, and appendices
                pages_to_process = min(total_pages, 500)
                
                # CRITICAL: Process in batches to avoid memory exhaustion
                # Each batch processes 50 pages to keep memory usage bounded
                BATCH_SIZE = 50
                page_results = {}
                
                # Progress tracking for large documents
                print(f"Processing {pages_to_process} pages with GPT-5.1 Vision extraction (parallel mode: {MAX_PARALLEL_PAGES} concurrent)...")
                print(f"Estimated cost: ${pages_to_process * 0.00096:.3f} | Time: ~{pages_to_process * 2 / MAX_PARALLEL_PAGES:.0f} seconds")
                
                for batch_start in range(1, pages_to_process + 1, BATCH_SIZE):
                    batch_end = min(batch_start + BATCH_SIZE - 1, pages_to_process)
                    batch_num = (batch_start - 1) // BATCH_SIZE + 1
                    total_batches = (pages_to_process + BATCH_SIZE - 1) // BATCH_SIZE
                    
                    print(f"Processing batch {batch_num}/{total_batches} (pages {batch_start}-{batch_end}) in parallel...")
                    
                    batch_images = convert_from_bytes(
                        file_bytes, 
                        dpi=200, 
                        fmt='jpeg', 
                        first_page=batch_start, 
                        last_page=batch_end
                    )
                    
                    page_tasks = []
                    for i, img in enumerate(batch_images):
                        page_num = batch_start + i
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='JPEG', quality=90)
                        img_bytes = img_byte_arr.getvalue()
                        page_tasks.append((img_bytes, page_num))
                    
                    with ThreadPoolExecutor(max_workers=MAX_PARALLEL_PAGES) as executor:
                        futures = {executor.submit(DocumentExtractor._process_single_page, task): task[1] for task in page_tasks}
                        for future in as_completed(futures):
                            page_num, page_text = future.result()
                            if page_text:
                                page_results[page_num] = page_text
                    
                    del batch_images
                
                vlm_text = "".join(page_results[pn] for pn in sorted(page_results.keys()))
                print(f"✓ Completed processing {pages_to_process} pages")
                
                if vlm_text.strip():
                    return vlm_text.strip()
                    
            except Exception as vlm_error:
                # VLM failed, try PyPDF2 as fallback
                pass
            
            # Fallback: Try standard text extraction if VLM completely failed
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                return text.strip()
            else:
                return "PDF text extraction incomplete - VLM and PyPDF2 both failed"
            
        except Exception as e:
            return f"Error extracting PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        try:
            doc = Document(io.BytesIO(file_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            return f"Error extracting DOCX: {str(e)}"
    
    @staticmethod
    def extract_text_from_xlsx(file_bytes: bytes) -> str:
        try:
            wb = load_workbook(filename=io.BytesIO(file_bytes), read_only=True)
            text = ""
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
            return text.strip()
        except Exception as e:
            return f"Error extracting XLSX: {str(e)}"
    
    @staticmethod
    def extract_text_from_image(file_bytes: bytes) -> str:
        """Enhanced image text extraction using VLM for better accuracy"""
        try:
            # Convert image to JPEG if needed
            image = Image.open(io.BytesIO(file_bytes))
            
            # Convert to RGB if needed (for RGBA images)
            if image.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
            
            # Convert to bytes
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='JPEG', quality=95)
            img_bytes = img_byte_arr.getvalue()
            
            # Use VLM for extraction
            vlm_text = DocumentExtractor._extract_text_with_vision(img_bytes, "image")
            
            # Check if VLM succeeded (not an error message)
            if not (vlm_text.startswith("VLM extraction failed") or vlm_text.startswith("VLM extraction unavailable")):
                return vlm_text
            
            # Fallback to pytesseract with enhanced configuration if VLM fails or is unavailable
            # Enhanced OCR configuration for better accuracy on technical documents
            custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
            # OEM 3: Default, based on what is available (LSTM + legacy)
            # PSM 6: Assume uniform block of text
            # preserve_interword_spaces: Better handling of tables and structured data
            text = pytesseract.image_to_string(image, config=custom_config)
            return text.strip() if text.strip() else f"Image text extraction incomplete"
            
        except Exception as e:
            return f"Error extracting text from image: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            return text.strip()
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    @staticmethod
    def is_drill_database(file_name: str, file_bytes: bytes) -> bool:
        """Check if file appears to be a drill database"""
        file_ext = file_name.lower().split('.')[-1]
        
        if file_ext not in ['csv', 'xlsx', 'xls']:
            return False
        
        keywords = ['drill', 'hole', 'collar', 'assay', 'sample', 'from_depth', 'to_depth', 
                   'dhid', 'hole_id', 'holeid', 'downhole', 'survey', 'azimuth', 'dip']
        
        file_name_lower = file_name.lower()
        return any(keyword in file_name_lower for keyword in keywords)
    
    @staticmethod
    def process_drill_database(file_name: str, file_bytes: bytes) -> Dict[str, Any]:
        """Process drill database and perform QAQC analysis"""
        try:
            parsed_data = DrillQAQCAnalyzer.parse_drill_database(file_bytes, file_name)
            
            if 'error' in parsed_data:
                return {
                    'file_name': file_name,
                    'file_type': 'drill_database',
                    'success': False,
                    'error': parsed_data['error']
                }
            
            qaqc_results = DrillQAQCAnalyzer.perform_full_analysis(parsed_data)
            
            return {
                'file_name': file_name,
                'file_type': 'drill_database',
                'success': True,
                'is_drill_database': True,
                'qaqc_results': qaqc_results,
                'text': qaqc_results.get('qaqc_report', 'QAQC analysis completed'),
                'qaqc_score': qaqc_results.get('qaqc_score', 0),
                'qaqc_rationale': qaqc_results.get('qaqc_rationale', '')
            }
        
        except Exception as e:
            return {
                'file_name': file_name,
                'file_type': 'drill_database',
                'success': False,
                'error': f'Error processing drill database: {str(e)}'
            }
    
    @staticmethod
    def extract_text(file_name: str, file_bytes: bytes) -> Dict[str, str]:
        file_ext = file_name.lower().split('.')[-1]
        
        if DocumentExtractor.is_drill_database(file_name, file_bytes):
            return DocumentExtractor.process_drill_database(file_name, file_bytes)
        
        extractors = {
            'pdf': DocumentExtractor.extract_text_from_pdf,
            'docx': DocumentExtractor.extract_text_from_docx,
            'doc': DocumentExtractor.extract_text_from_docx,
            'xlsx': DocumentExtractor.extract_text_from_xlsx,
            'xls': DocumentExtractor.extract_text_from_xlsx,
            'csv': DocumentExtractor.extract_text_from_txt,
            'txt': DocumentExtractor.extract_text_from_txt,
            'jpg': DocumentExtractor.extract_text_from_image,
            'jpeg': DocumentExtractor.extract_text_from_image,
            'png': DocumentExtractor.extract_text_from_image,
        }
        
        extractor = extractors.get(file_ext)
        
        if extractor:
            extracted_text = extractor(file_bytes)
            return {
                'file_name': file_name,
                'file_type': file_ext,
                'text': extracted_text,
                'success': not extracted_text.startswith('Error'),
                'is_drill_database': False
            }
        else:
            return {
                'file_name': file_name,
                'file_type': file_ext,
                'text': '',
                'success': False,
                'error': f'Unsupported file type: {file_ext}',
                'is_drill_database': False
            }
